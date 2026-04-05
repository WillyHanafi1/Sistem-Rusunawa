from pathlib import Path
from docx import Document

templates_dir = Path(r"D:\ProjectAI\Sistem-Rusunawa\rusun-backend\app\templates")
templates_dir.mkdir(parents=True, exist_ok=True)

templates = [
    "template_skrd.docx",
    "template_strd.docx",
    "template_teguran1.docx",
    "template_teguran2.docx",
    "template_teguran3.docx"
]

for template_name in templates:
    doc = Document()
    doc.add_heading(template_name.replace("template_", "").replace(".docx", "").upper(), 0)
    doc.add_paragraph(f"Ini adalah dokumen dummy untuk {template_name}")
    
    # Generic placeholders based on typical usage to satisfy basic templating rendering
    doc.add_paragraph("Penyewa: {{ name }}")
    doc.add_paragraph("Unit: {{ unit_number }}")
    doc.add_paragraph("Bulan: {{ month }}")
    doc.add_paragraph("Tahun: {{ year }}")
    doc.add_paragraph("Nominal: {{ amount }}")
    doc.add_paragraph("Tanggal Cetak: {{ tanggal_cetak }}")
    doc.add_paragraph("Nama Kepala UPTD: {{ nama_kepala_uptd }}")
    
    doc_path = templates_dir / template_name
    
    # Python-docx overwrites existing files easily
    print(f"Creating {doc_path}...")
    doc.save(str(doc_path))

print("Done! Templates generated.")
