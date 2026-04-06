import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_ba_wawancara(path):
    doc = Document()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("PEMERINTAH DAERAH KOTA CIMAHI\nDINAS PERUMAHAN DAN KAWASAN PERMUKIMAN\nUPTD RUMAH SUSUN SEDERHANA SEWA (RUSUNAWA)")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph("Jl. At-taqwa Kelurahan Cigugur Tengah Kota Cimahi 40522\nTelp. (022) 6653680").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BERITA ACARA WAWANCARA PENGHUNI\nNOMOR: {{ sk_nomor|default('.../2026') }}")
    run.bold = True
    
    p = doc.add_paragraph("\nPada hari ini {{ hari }}, tanggal {{ tanggal_indo }}, telah dilaksanakan wawancara terhadap:")
    doc.add_paragraph("Nama: {{ nama_penyewa }}")
    doc.add_paragraph("NIK: {{ nik }}")
    doc.add_paragraph("Pekerjaan: {{ pekerjaan }}")
    doc.add_paragraph("Alamat: {{ alamat_asal }}")
    
    doc.add_paragraph("\nHasil Wawancara: DITERIMA / TIDAK DITERIMA")
    doc.add_paragraph("Kamar yang ditetapkan: {{ nomor_kamar }} (Lantai {{ lantai }}, Blok {{ gedung }})")
    doc.add_paragraph("Tarif Sewa: {{ harga_sewa }} / bulan")
    
    doc.add_paragraph("\nDemikian Berita Acara ini dibuat untuk dipergunakan sebagaimana mestinya.")
    
    doc.add_paragraph("\n\nCimahi, {{ tanggal_cetak }}")
    doc.add_paragraph("Kepala UPTD Rusunawa,\n\n\n\n{{ nama_kepala_uptd }}\nNIP. {{ nip_kepala_uptd }}")
    doc.save(path)

def create_perjanjian_sewa(path):
    doc = Document()
    title = doc.add_paragraph("SURAT PERJANJIAN SEWA MENYEWA UNIT HUNIAN")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    
    doc.add_paragraph("\nAntara:")
    doc.add_paragraph("1. NAMA: {{ nama_kepala_uptd }} (Selaku Pihak Pertama)")
    doc.add_paragraph("2. NAMA: {{ nama_penyewa }} (Selaku Pihak Kedua)")
    
    doc.add_paragraph("\nPasal 1: OBJEK SEWA")
    doc.add_paragraph("Pihak Pertama menyewakan kepada Pihak Kedua unit hunian {{ nomor_kamar }} di Rusunawa {{ rusunawa }}.")
    
    doc.add_paragraph("\nPasal 2: JANGKA WAKTU")
    doc.add_paragraph("Mulai: {{ tanggal_mulai }} s/d {{ tanggal_selesai }}")
    
    doc.add_paragraph("\nPasal 3: HARGA & JAMINAN")
    doc.add_paragraph("Harga Sewa: {{ harga_sewa }}")
    doc.add_paragraph("Uang Jaminan (2x Sewa): {{ deposit }}")
    
    doc.add_paragraph("\nPIHAK PERTAMA\t\t\tPIHAK KEDUA")
    doc.add_paragraph("\n\n( {{ nama_kepala_uptd }} )\t\t\t( {{ nama_penyewa }} )")
    doc.save(path)

def create_ba_serah_terima(path):
    doc = Document()
    doc.add_heading("BERITA ACARA SERAH TERIMA UNIT & KUNCI", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\nSudah diserahkan unit hunian ke:")
    doc.add_paragraph("Nama: {{ nama_penyewa }}")
    doc.add_paragraph("Unit: {{ nomor_kamar }}")
    doc.add_paragraph("Gedung: {{ gedung }}, Lantai: {{ lantai }}")
    doc.add_paragraph("\nKunci yang diserahkan: 1 Set Kunci")
    doc.add_paragraph("Jam Masuk: {{ jam_masuk }}")
    doc.save(path)

def create_form_pendaftaran(path, type_name):
    doc = Document()
    doc.add_heading(f"FORM PENDAFTARAN MBR ({type_name})", 0)
    doc.add_paragraph("Nama Lengkap: {{ nama_penyewa }}")
    doc.add_paragraph("NIK: {{ nik }}")
    doc.add_paragraph("Pekerjaan: {{ pekerjaan }}")
    doc.add_paragraph("Rusunawa Target: {{ rusunawa }}")
    doc.save(path)

def create_ba_checkout(path):
    doc = Document()
    doc.add_heading("BERITA ACARA PENGOSONGAN UNIT (CHECKOUT)", 0)
    doc.add_paragraph("Nama: {{ nama_penyewa }}")
    doc.add_paragraph("Unit: {{ nomor_kamar }}")
    doc.add_paragraph("Tanggal Keluar: {{ tanggal_cetak }}")
    doc.save(path)

def create_surat_jalan(path):
    doc = Document()
    doc.add_heading("SURAT IZIN MASUK BARANG (SURAT JALAN)", 0)
    doc.add_paragraph("Memberikan izin kepada {{ nama_penyewa }} untuk memasukkan barang-barang pindahan ke unit {{ nomor_kamar }}.")
    doc.save(path)

def create_invoice_template(path, title):
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Nomor: {{ nomor_dokumen }}")
    doc.add_paragraph("Nama: {{ nama_penyewa }}")
    doc.add_paragraph("Unit: {{ nomor_kamar }}")
    doc.add_paragraph("Bulan: {{ bulan_indo }} {{ tahun }}")
    doc.add_paragraph("\nDetail Tagihan:")
    doc.add_paragraph("- Sewa Kamar: {{ sewa_kamar }}")
    doc.add_paragraph("- Air: {{ air }}")
    doc.add_paragraph("- Listrik: {{ listrik }}")
    doc.add_paragraph("- Kebersihan/Lainnya: {{ kebersihan }}")
    doc.add_paragraph("\nTOTAL: {{ total }}")
    doc.save(path)

def create_teguran_template(path, level):
    doc = Document()
    doc.add_heading(f"SURAT TEGURAN KE-{level}", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Kepada Yth. {{ nama_penyewa }}\ndi Unit {{ nomor_kamar }}")
    doc.add_paragraph("\nBerdasarkan catatan kami, Anda belum menyelesaikan kewajiban pembayaran untuk bulan {{ bulan_indo }} {{ tahun }}.")
    doc.add_paragraph("Mohon segera melakukan pembayaran sebesar {{ total }}.")
    doc.add_paragraph("\nDemikian, mohon perhatiannya.")
    doc.add_paragraph("\nKepala UPTD,\n\n\n{{ nama_kepala_uptd }}")
    doc.save(path)

template_dir = r"d:\ProjectAI\Sistem-Rusunawa\rusun-backend\app\templates"
os.makedirs(template_dir, exist_ok=True)

create_ba_wawancara(os.path.join(template_dir, "ba_wawancara.docx"))
create_perjanjian_sewa(os.path.join(template_dir, "perjanjian_sewa.docx"))
create_ba_serah_terima(os.path.join(template_dir, "ba_serah_terima.docx"))
create_form_pendaftaran(os.path.join(template_dir, "form_pendaftaran_karyawan.docx"), "KARYAWAN")
create_form_pendaftaran(os.path.join(template_dir, "form_pendaftaran_wiraswasta.docx"), "WIRASWASTA")
create_ba_checkout(os.path.join(template_dir, "ba_checkout.docx"))
create_surat_jalan(os.path.join(template_dir, "surat_jalan_masuk.docx"))

# Extra system templates
create_invoice_template(os.path.join(template_dir, "template_skrd.docx"), "SURAT KETETAPAN RETRIBUSI DAERAH (SKRD)")
create_invoice_template(os.path.join(template_dir, "template_strd.docx"), "SURAT TAGIHAN RETRIBUSI DAERAH (STRD)")
create_teguran_template(os.path.join(template_dir, "template_teguran_1.docx"), "1")
create_teguran_template(os.path.join(template_dir, "template_teguran_2.docx"), "2")
create_teguran_template(os.path.join(template_dir, "template_teguran_3.docx"), "3")

print("Semua template .docx berhasil dibuat di " + template_dir)
